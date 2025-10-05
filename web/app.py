"""Web application for Voice-Driven Brainstorming Assistant."""
import os
import sys
from pathlib import Path
from flask import Flask, render_template, request, jsonify, session
from flask_socketio import SocketIO, emit
import numpy as np
import base64
import json
from datetime import datetime
import uuid

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from brain.model import BrainstormSession, IdeaSource
from brain.organizer import Organizer
from brain.assistant import BrainstormAssistant
from stt.whisper_local import WhisperLocalSTT
from llm.openai_client import OpenAIClient
from storage.files import FileStorage
from storage.exporters import export_session
from utils.config import Config
from utils.logging import get_logger

logger = get_logger('web')

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.urandom(24)
socketio = SocketIO(app, cors_allowed_origins="*")

# Global state
sessions = {}  # session_id -> BrainstormSession
organizers = {}  # session_id -> Organizer
assistants = {}  # session_id -> BrainstormAssistant
storages = {}  # session_id -> FileStorage

# Initialize backends
config = Config()
stt_backend = None
llm_backend = None

def init_backends():
    """Initialize STT and LLM backends."""
    global stt_backend, llm_backend
    
    # Initialize STT
    try:
        stt_backend = WhisperLocalSTT(
            model_size=config.get('stt.whisper_model', 'base'),
            sample_rate=16000
        )
        logger.info("STT backend initialized")
    except Exception as e:
        logger.error(f"Failed to initialize STT: {e}")
    
    # Initialize LLM
    try:
        if config.openai_api_key:
            llm_backend = OpenAIClient(
                api_key=config.openai_api_key,
                model=config.openai_model
            )
            logger.info("LLM backend initialized")
        else:
            logger.warning("No OpenAI API key configured")
    except Exception as e:
        logger.error(f"Failed to initialize LLM: {e}")

def get_or_create_session(session_id: str, project_name: str = "default"):
    """Get or create a brainstorming session."""
    if session_id not in sessions:
        # Create storage
        storage = FileStorage(Path("brainstorm") / project_name)
        
        # Load or create session
        if storage.exists():
            brainstorm_session = storage.load_session()
        else:
            brainstorm_session = BrainstormSession(project_name=project_name)
        
        # Create organizer
        organizer = Organizer(brainstorm_session)
        
        # Create assistant
        assistant = None
        if llm_backend:
            assistant = BrainstormAssistant(llm_backend, organizer)
        
        # Store
        sessions[session_id] = brainstorm_session
        organizers[session_id] = organizer
        assistants[session_id] = assistant
        storages[session_id] = storage
        
        logger.info(f"Created session: {session_id} for project: {project_name}")
    
    return sessions[session_id], organizers[session_id], assistants.get(session_id), storages[session_id]

@app.route('/')
def index():
    """Render main page."""
    return render_template('index.html')

@app.route('/api/session/create', methods=['POST'])
def create_session():
    """Create a new session."""
    data = request.json
    project_name = data.get('project_name', 'default')
    
    # Generate session ID
    session_id = os.urandom(16).hex()
    session['session_id'] = session_id
    
    # Create session
    brainstorm_session, organizer, assistant, storage = get_or_create_session(session_id, project_name)
    
    return jsonify({
        'success': True,
        'session_id': session_id,
        'project_name': brainstorm_session.project_name
    })

@app.route('/api/session/data', methods=['GET'])
def get_session_data():
    """Get current session data."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    brainstorm_session = sessions[session_id]
    organizer = organizers[session_id]
    
    # Get active ideas
    active_ideas = brainstorm_session.get_active_ideas()
    
    return jsonify({
        'success': True,
        'project_name': brainstorm_session.project_name,
        'ideas': [idea.to_dict() for idea in active_ideas],
        'clusters': [cluster.to_dict() for cluster in brainstorm_session.clusters],
        'actions': [action.to_dict() for action in brainstorm_session.actions],
        'transcript': [entry.to_dict() for entry in brainstorm_session.transcript[-20:]],  # Last 20
        'stats': {
            'total_ideas': len(active_ideas),
            'total_clusters': len(brainstorm_session.clusters),
            'total_actions': len([a for a in brainstorm_session.actions if not a.completed]),
            'key_ideas': len([i for i in active_ideas if i.promoted])
        }
    })

@app.route('/api/transcribe', methods=['POST'])
def transcribe_audio():
    """Transcribe audio data."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    if not stt_backend:
        return jsonify({'success': False, 'error': 'STT backend not available'})
    
    # Get audio data
    data = request.json
    audio_base64 = data.get('audio')
    
    if not audio_base64:
        return jsonify({'success': False, 'error': 'No audio data'})
    
    try:
        # Decode audio
        audio_bytes = base64.b64decode(audio_base64)
        audio_array = np.frombuffer(audio_bytes, dtype=np.int16)
        
        logger.info(f"Received audio: {len(audio_array)} samples, duration: {len(audio_array)/16000:.2f}s")
        
        # Check if audio is too short
        if len(audio_array) < 1600:  # Less than 0.1 seconds
            return jsonify({'success': False, 'error': 'Audio too short'})
        
        # Transcribe
        text = stt_backend.transcribe(audio_array)
        
        if not text:
            logger.warning("Transcription returned empty")
            return jsonify({'success': False, 'error': 'No speech detected'})
        
        # Add to session
        brainstorm_session, organizer, assistant, storage = get_or_create_session(session_id)
        
        # Add transcript
        organizer.add_transcript(text, "user")
        
        # Add idea
        idea = organizer.add_idea(text, source=IdeaSource.USER)
        
        # Get assistant response
        assistant_response = None
        if assistant:
            assistant_response = assistant.process_user_input(text)
            if assistant_response:
                organizer.add_transcript(assistant_response, "assistant")
        
        # Save
        storage.save_session(brainstorm_session)
        
        return jsonify({
            'success': True,
            'text': text,
            'idea': idea.to_dict(),
            'assistant_response': assistant_response
        })
        
    except Exception as e:
        logger.error(f"Transcription error: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/idea/tag', methods=['POST'])
def tag_idea():
    """Tag an idea."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    data = request.json
    idea_id = data.get('idea_id')
    tags = data.get('tags', [])
    
    organizer = organizers[session_id]
    success = organizer.tag_idea(idea_id, tags)
    
    if success:
        storage = storages[session_id]
        storage.save_session(sessions[session_id])
    
    return jsonify({'success': success})

@app.route('/api/idea/promote', methods=['POST'])
def promote_idea():
    """Promote an idea."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    data = request.json
    idea_id = data.get('idea_id')
    
    organizer = organizers[session_id]
    success = organizer.promote_idea(idea_id)
    
    if success:
        storage = storages[session_id]
        storage.save_session(sessions[session_id])
    
    return jsonify({'success': success})

@app.route('/api/idea/delete', methods=['POST'])
def delete_idea():
    """Delete an idea."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    data = request.json
    idea_id = data.get('idea_id')
    
    organizer = organizers[session_id]
    success = organizer.delete_idea(idea_id, soft=True)
    
    if success:
        storage = storages[session_id]
        storage.save_session(sessions[session_id])
    
    return jsonify({'success': success})

@app.route('/api/action/add', methods=['POST'])
def add_action():
    """Add an action item."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    data = request.json
    text = data.get('text')
    
    organizer = organizers[session_id]
    action = organizer.add_action(text)
    
    storage = storages[session_id]
    storage.save_session(sessions[session_id])
    
    return jsonify({'success': True, 'action': action.to_dict()})

@app.route('/api/export/<format>', methods=['GET'])
def export(format):
    """Export session."""
    session_id = session.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify({'success': False, 'error': 'No active session'})
    
    brainstorm_session = sessions[session_id]
    storage = storages[session_id]
    
    results = export_session(brainstorm_session, storage.base_dir, [format])
    
    if results and format in results:
        return jsonify({'success': True, 'file': str(results[format])})
    else:
        return jsonify({'success': False, 'error': 'Export failed'})

# ============================================
# DOCUMENT MODE ENDPOINTS
# ============================================

# Document storage directory (separate from brainstorm sessions)
DOCUMENTS_DIR = Path("documents")
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

def get_document_path(doc_id):
    """Get path to document file."""
    return DOCUMENTS_DIR / f"{doc_id}.json"

def load_document_from_file(doc_id):
    """Load document from file."""
    doc_path = get_document_path(doc_id)
    if doc_path.exists():
        with open(doc_path, 'r') as f:
            return json.load(f)
    return None

def save_document_to_file(document):
    """Save document to file."""
    doc_path = get_document_path(document['id'])
    try:
        with open(doc_path, 'w', encoding='utf-8') as f:
            json.dump(document, f, indent=2, ensure_ascii=False)
        logger.info(f"Saved document to: {doc_path}")
    except Exception as e:
        logger.error(f"Failed to write document file {doc_path}: {e}")
        raise

@app.route('/api/documents/list', methods=['GET'])
def list_documents():
    """List all documents."""
    try:
        documents = []
        for doc_file in DOCUMENTS_DIR.glob("*.json"):
            try:
                with open(doc_file, 'r') as f:
                    doc = json.load(f)
                    documents.append({
                        'id': doc['id'],
                        'title': doc['title'],
                        'word_count': doc.get('word_count', 0),
                        'updated_at': doc.get('updated_at', doc.get('created_at'))
                    })
            except Exception as e:
                logger.error(f"Failed to load document {doc_file}: {e}")
        
        # Sort by updated_at descending
        documents.sort(key=lambda x: x['updated_at'], reverse=True)
        
        return jsonify({'success': True, 'documents': documents})
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/documents/create', methods=['POST'])
def create_document():
    """Create a new document."""
    try:
        data = request.json
        title = data.get('title', 'Untitled Document')
        content = data.get('content', '')
        
        doc_id = str(uuid.uuid4())
        now = datetime.now().isoformat()
        
        # Calculate word count properly
        word_count = len(content.strip().split()) if content.strip() else 0
        
        document = {
            'id': doc_id,
            'title': title,
            'content': content,
            'word_count': word_count,
            'created_at': now,
            'updated_at': now
        }
        
        logger.info(f"Creating document: {doc_id}, title: {title}")
        save_document_to_file(document)
        logger.info(f"Document created successfully: {doc_id}")
        
        return jsonify({'success': True, 'document': document})
    except Exception as e:
        logger.error(f"Failed to create document: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/documents/get/<doc_id>', methods=['GET'])
def get_document(doc_id):
    """Get a specific document."""
    try:
        document = load_document_from_file(doc_id)
        if document:
            return jsonify({'success': True, 'document': document})
        else:
            return jsonify({'success': False, 'error': 'Document not found'})
    except Exception as e:
        logger.error(f"Failed to get document: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/documents/save/<doc_id>', methods=['POST'])
def save_document(doc_id):
    """Save document changes."""
    try:
        data = request.json
        document = load_document_from_file(doc_id)
        
        if not document:
            return jsonify({'success': False, 'error': 'Document not found'})
        
        # Update document
        document['title'] = data.get('title', document['title'])
        document['content'] = data.get('content', document['content'])
        
        # Calculate word count properly
        content = document['content']
        document['word_count'] = len(content.strip().split()) if content.strip() else 0
        document['updated_at'] = datetime.now().isoformat()
        
        save_document_to_file(document)
        
        logger.info(f"Document saved: {document['id']}, title: {document['title']}, words: {document['word_count']}")
        
        return jsonify({'success': True, 'document': document})
    except Exception as e:
        logger.error(f"Failed to save document: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/documents/delete/<doc_id>', methods=['POST'])
def delete_document(doc_id):
    """Delete a document."""
    try:
        doc_path = get_document_path(doc_id)
        if doc_path.exists():
            doc_path.unlink()
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'Document not found'})
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/documents/transcribe', methods=['POST'])
def transcribe_for_document():
    """Transcribe audio for document mode (doesn't save to brainstorm session)."""
    if not stt_backend:
        return jsonify({'success': False, 'error': 'STT backend not available'})
    
    # Get audio data
    data = request.json
    audio_base64 = data.get('audio')
    
    if not audio_base64:
        return jsonify({'success': False, 'error': 'No audio data'})
    
    try:
        # Decode audio
        audio_bytes = base64.b64decode(audio_base64)
        audio_array = np.frombuffer(audio_bytes, dtype=np.int16)
        
        logger.info(f"Document transcription - Received audio: {len(audio_array)} samples, duration: {len(audio_array)/16000:.2f}s")
        
        # Check if audio is too short
        if len(audio_array) < 1600:  # Less than 0.1 seconds
            return jsonify({'success': False, 'error': 'Audio too short'})
        
        # Transcribe
        text = stt_backend.transcribe(audio_array)
        
        if not text:
            logger.warning("Document transcription returned empty")
            return jsonify({'success': False, 'error': 'No speech detected'})
        
        logger.info(f"Document transcription successful: {len(text)} characters")
        
        # Return just the transcribed text (no brainstorm session saving)
        return jsonify({
            'success': True,
            'text': text
        })
        
    except Exception as e:
        logger.error(f"Document transcription error: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/documents/export/<doc_id>/<format>', methods=['GET'])
def export_document(doc_id, format):
    """Export a document to various formats."""
    try:
        document = load_document_from_file(doc_id)
        if not document:
            return jsonify({'success': False, 'error': 'Document not found'})
        
        title = document['title']
        content = document['content']
        
        # Create exports directory inside documents folder
        exports_dir = DOCUMENTS_DIR / 'exports'
        exports_dir.mkdir(exist_ok=True)
        
        # Sanitize filename
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
        if not safe_title:
            safe_title = doc_id[:8]
        
        if format == 'md':
            # Export as Markdown
            filename = f"{safe_title}.md"
            filepath = exports_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"# {title}\n\n")
                f.write(f"*Created: {document['created_at']}*\n\n")
                f.write(f"*Last Updated: {document['updated_at']}*\n\n")
                f.write("---\n\n")
                f.write(content)
            
            logger.info(f"Exported document to Markdown: {filepath}")
            return jsonify({'success': True, 'file': str(filepath), 'filename': filename})
        
        elif format == 'txt':
            # Export as plain text
            filename = f"{safe_title}.txt"
            filepath = exports_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"{title}\n")
                f.write("=" * len(title) + "\n\n")
                f.write(content)
            
            logger.info(f"Exported document to TXT: {filepath}")
            return jsonify({'success': True, 'file': str(filepath), 'filename': filename})
        
        elif format == 'docx':
            # Export as DOCX
            try:
                from docx import Document as DocxDocument
                from docx.shared import Pt, Inches
                
                filename = f"{safe_title}.docx"
                filepath = exports_dir / filename
                
                doc = DocxDocument()
                
                # Add title
                title_para = doc.add_heading(title, level=1)
                
                # Add metadata
                meta_para = doc.add_paragraph()
                meta_para.add_run(f"Created: {document['created_at']}\n").italic = True
                meta_para.add_run(f"Last Updated: {document['updated_at']}\n").italic = True
                meta_para.add_run(f"Word Count: {document['word_count']}").italic = True
                
                doc.add_paragraph()  # Spacing
                
                # Add content
                for paragraph in content.split('\n\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                
                doc.save(str(filepath))
                
                logger.info(f"Exported document to DOCX: {filepath}")
                return jsonify({'success': True, 'file': str(filepath), 'filename': filename})
            
            except ImportError:
                logger.error("python-docx not installed")
                return jsonify({'success': False, 'error': 'DOCX export requires python-docx package'})
        
        elif format == 'pdf':
            # Export as PDF (requires reportlab)
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                
                filename = f"{safe_title}.pdf"
                filepath = exports_dir / filename
                
                doc = SimpleDocTemplate(str(filepath), pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=24,
                    spaceAfter=30,
                )
                story.append(Paragraph(title, title_style))
                
                # Metadata
                meta_text = f"<i>Created: {document['created_at']}<br/>Last Updated: {document['updated_at']}<br/>Word Count: {document['word_count']}</i>"
                story.append(Paragraph(meta_text, styles['Normal']))
                story.append(Spacer(1, 0.3*inch))
                
                # Content
                for paragraph in content.split('\n\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph.strip(), styles['Normal']))
                        story.append(Spacer(1, 0.2*inch))
                
                doc.build(story)
                
                logger.info(f"Exported document to PDF: {filepath}")
                return jsonify({'success': True, 'file': str(filepath), 'filename': filename})
            
            except ImportError:
                logger.error("reportlab not installed")
                return jsonify({'success': False, 'error': 'PDF export requires reportlab package'})
        
        else:
            return jsonify({'success': False, 'error': f'Unsupported format: {format}'})
    
    except Exception as e:
        logger.error(f"Failed to export document: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    logger.info("Starting web application...")
    init_backends()
    # Use 127.0.0.1 (localhost) for microphone access - browsers treat it as secure origin
    socketio.run(app, host='127.0.0.1', port=5000, debug=True)

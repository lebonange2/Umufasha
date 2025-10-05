"""Export brainstorming sessions to various formats."""
from pathlib import Path
from typing import Optional
from datetime import datetime
import csv
from brain.model import BrainstormSession
from utils.logging import get_logger

logger = get_logger('storage.exporters')


class MarkdownExporter:
    """Export session to Markdown format."""
    
    @staticmethod
    def export(session: BrainstormSession, filepath: Path, 
               include_transcript: bool = True, include_timestamps: bool = False) -> bool:
        """Export session to Markdown.
        
        Args:
            session: Session to export
            filepath: Output file path
            include_transcript: Include full transcript
            include_timestamps: Include timestamps
            
        Returns:
            True if successful
        """
        try:
            lines = []
            
            # Header
            lines.append(f"# {session.project_name}")
            lines.append("")
            lines.append(f"**Created**: {session.created_at.strftime('%Y-%m-%d %H:%M')}")
            lines.append(f"**Updated**: {session.updated_at.strftime('%Y-%m-%d %H:%M')}")
            lines.append("")
            lines.append("---")
            lines.append("")
            
            # Key Ideas
            key_ideas = session.get_key_ideas()
            if key_ideas:
                lines.append("## 🌟 Key Ideas")
                lines.append("")
                for idea in key_ideas:
                    tags_str = f" `{', '.join(idea.tags)}`" if idea.tags else ""
                    timestamp_str = f" *({idea.timestamp.strftime('%H:%M')})*" if include_timestamps else ""
                    lines.append(f"- **{idea.text}**{tags_str}{timestamp_str}")
                lines.append("")
            
            # All Active Ideas
            active_ideas = [i for i in session.get_active_ideas() if not i.promoted]
            if active_ideas:
                lines.append("## 💡 Ideas")
                lines.append("")
                for idea in active_ideas:
                    tags_str = f" `{', '.join(idea.tags)}`" if idea.tags else ""
                    timestamp_str = f" *({idea.timestamp.strftime('%H:%M')})*" if include_timestamps else ""
                    source_icon = "🤖" if idea.source.value == "assistant" else "👤"
                    lines.append(f"- {source_icon} {idea.text}{tags_str}{timestamp_str}")
                lines.append("")
            
            # Clusters
            if session.clusters:
                lines.append("## 🗂️ Clusters")
                lines.append("")
                for cluster in session.clusters:
                    tags_str = f" `{', '.join(cluster.tags)}`" if cluster.tags else ""
                    lines.append(f"### {cluster.name}{tags_str}")
                    if cluster.description:
                        lines.append(f"*{cluster.description}*")
                        lines.append("")
                    
                    # List ideas in cluster
                    for idea_id in cluster.idea_ids:
                        idea = session.get_idea(idea_id)
                        if idea and idea.merged_into is None:
                            lines.append(f"- {idea.text}")
                    lines.append("")
            
            # Action Items
            if session.actions:
                lines.append("## ✅ Action Items")
                lines.append("")
                
                # Active actions
                active_actions = [a for a in session.actions if not a.completed]
                if active_actions:
                    lines.append("### To Do")
                    lines.append("")
                    for action in active_actions:
                        priority_emoji = {"urgent": "🔴", "high": "🟠", "medium": "🟡", "low": "🟢"}.get(action.priority.value, "⚪")
                        due_str = f" (Due: {action.due_date.strftime('%Y-%m-%d')})" if action.due_date else ""
                        lines.append(f"- [ ] {priority_emoji} {action.text}{due_str}")
                    lines.append("")
                
                # Completed actions
                completed_actions = [a for a in session.actions if a.completed]
                if completed_actions:
                    lines.append("### Completed")
                    lines.append("")
                    for action in completed_actions:
                        lines.append(f"- [x] {action.text}")
                    lines.append("")
            
            # Summaries
            if session.summaries:
                lines.append("## 📝 Summaries")
                lines.append("")
                for summary in session.summaries:
                    lines.append(f"### {summary.scope.title()}")
                    lines.append("")
                    lines.append(summary.text)
                    lines.append("")
            
            # Transcript
            if include_transcript and session.transcript:
                lines.append("## 📜 Transcript")
                lines.append("")
                lines.append("<details>")
                lines.append("<summary>Click to expand</summary>")
                lines.append("")
                
                for entry in session.transcript:
                    speaker_icon = "🤖 **Assistant**" if entry.speaker == "assistant" else "👤 **User**"
                    timestamp_str = entry.timestamp.strftime('%H:%M:%S')
                    lines.append(f"**[{timestamp_str}]** {speaker_icon}")
                    lines.append(f"> {entry.text}")
                    lines.append("")
                
                lines.append("</details>")
                lines.append("")
            
            # Statistics
            lines.append("---")
            lines.append("")
            lines.append("## 📊 Statistics")
            lines.append("")
            lines.append(f"- **Total Ideas**: {len(session.get_active_ideas())}")
            lines.append(f"- **Key Ideas**: {len(key_ideas)}")
            lines.append(f"- **Clusters**: {len(session.clusters)}")
            lines.append(f"- **Action Items**: {len([a for a in session.actions if not a.completed])}")
            lines.append(f"- **Completed Actions**: {len([a for a in session.actions if a.completed])}")
            lines.append(f"- **Transcript Entries**: {len(session.transcript)}")
            
            # Write to file
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write('\n'.join(lines))
            
            logger.info(f"Exported to Markdown: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to export Markdown: {e}")
            return False


class CSVExporter:
    """Export session to CSV format."""
    
    @staticmethod
    def export(session: BrainstormSession, filepath: Path) -> bool:
        """Export ideas to CSV.
        
        Args:
            session: Session to export
            filepath: Output file path
            
        Returns:
            True if successful
        """
        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Header
                writer.writerow(['ID', 'Text', 'Tags', 'Source', 'Promoted', 'Score', 'Timestamp'])
                
                # Ideas
                for idea in session.get_active_ideas():
                    writer.writerow([
                        idea.id,
                        idea.text,
                        ', '.join(idea.tags),
                        idea.source.value,
                        'Yes' if idea.promoted else 'No',
                        idea.score,
                        idea.timestamp.isoformat()
                    ])
            
            logger.info(f"Exported to CSV: {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to export CSV: {e}")
            return False


class DOCXExporter:
    """Export session to DOCX format."""
    
    @staticmethod
    def export(session: BrainstormSession, filepath: Path) -> bool:
        """Export session to DOCX.
        
        Args:
            session: Session to export
            filepath: Output file path
            
        Returns:
            True if successful
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # Title
            title = doc.add_heading(session.project_name, 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Metadata
            p = doc.add_paragraph()
            p.add_run(f"Created: {session.created_at.strftime('%Y-%m-%d %H:%M')}\n").italic = True
            p.add_run(f"Updated: {session.updated_at.strftime('%Y-%m-%d %H:%M')}").italic = True
            
            doc.add_page_break()
            
            # Key Ideas
            key_ideas = session.get_key_ideas()
            if key_ideas:
                doc.add_heading('Key Ideas', 1)
                for idea in key_ideas:
                    p = doc.add_paragraph(idea.text, style='List Bullet')
                    if idea.tags:
                        p.add_run(f" [{', '.join(idea.tags)}]").italic = True
            
            # All Ideas
            active_ideas = [i for i in session.get_active_ideas() if not i.promoted]
            if active_ideas:
                doc.add_heading('Ideas', 1)
                for idea in active_ideas:
                    p = doc.add_paragraph(idea.text, style='List Bullet')
                    if idea.tags:
                        p.add_run(f" [{', '.join(idea.tags)}]").italic = True
            
            # Clusters
            if session.clusters:
                doc.add_heading('Clusters', 1)
                for cluster in session.clusters:
                    doc.add_heading(cluster.name, 2)
                    if cluster.description:
                        doc.add_paragraph(cluster.description).italic = True
                    
                    for idea_id in cluster.idea_ids:
                        idea = session.get_idea(idea_id)
                        if idea and idea.merged_into is None:
                            doc.add_paragraph(idea.text, style='List Bullet 2')
            
            # Action Items
            if session.actions:
                doc.add_heading('Action Items', 1)
                
                active_actions = [a for a in session.actions if not a.completed]
                if active_actions:
                    doc.add_heading('To Do', 2)
                    for action in active_actions:
                        p = doc.add_paragraph(action.text, style='List Bullet')
                        p.add_run(f" [{action.priority.value}]").bold = True
                
                completed_actions = [a for a in session.actions if a.completed]
                if completed_actions:
                    doc.add_heading('Completed', 2)
                    for action in completed_actions:
                        doc.add_paragraph(action.text, style='List Bullet')
            
            # Summaries
            if session.summaries:
                doc.add_heading('Summaries', 1)
                for summary in session.summaries:
                    doc.add_heading(summary.scope.title(), 2)
                    doc.add_paragraph(summary.text)
            
            # Save
            doc.save(filepath)
            
            logger.info(f"Exported to DOCX: {filepath}")
            return True
            
        except ImportError:
            logger.error("python-docx not installed")
            return False
        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            return False


def export_session(session: BrainstormSession, base_dir: Path, 
                   formats: list = ['md', 'json', 'csv']) -> dict:
    """Export session to multiple formats.
    
    Args:
        session: Session to export
        base_dir: Base directory for exports
        formats: List of formats to export ('md', 'json', 'csv', 'docx')
        
    Returns:
        Dictionary of format -> filepath for successful exports
    """
    results = {}
    
    if 'md' in formats:
        md_path = base_dir / "notes.md"
        if MarkdownExporter.export(session, md_path):
            results['md'] = md_path
    
    if 'csv' in formats:
        csv_path = base_dir / "ideas.csv"
        if CSVExporter.export(session, csv_path):
            results['csv'] = csv_path
    
    if 'docx' in formats:
        docx_path = base_dir / "notes.docx"
        if DOCXExporter.export(session, docx_path):
            results['docx'] = docx_path
    
    return results
